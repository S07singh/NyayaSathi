"""
NyayaSathi AI — Document Parser
================================
Extracts clean text from uploaded PDF and DOCX files.

Supported formats:
  • PDF  → pdfplumber  (handles scanned-text PDFs)
  • DOCX → python-docx (reads paragraphs + tables)
"""

from __future__ import annotations

import io
from pathlib import Path
from typing import Union

import pdfplumber
from docx import Document as DocxDocument

from utils.logger import get_logger
from utils.text_cleaner import clean_text

logger = get_logger(__name__)


def parse_document(file: Union[str, Path, io.BytesIO], filename: str = "") -> str:
    """Auto-detect format and return cleaned text.

    Parameters
    ----------
    file : str | Path | BytesIO
        File path or in-memory buffer (e.g. from Streamlit uploader).
    filename : str
        Original filename — used for extension detection when *file*
        is a BytesIO buffer.

    Returns
    -------
    str
        Cleaned document text.

    Raises
    ------
    ValueError
        If the file format is unsupported.
    """
    ext = _detect_extension(file, filename)
    logger.info("Parsing document (format: %s)", ext)

    if ext == ".pdf":
        raw = parse_pdf(file)
    elif ext in (".docx", ".doc"):
        raw = parse_docx(file)
    else:
        raise ValueError(f"Unsupported file format: '{ext}'. Upload a PDF or DOCX.")

    cleaned = clean_text(raw)
    logger.info("Extracted %d characters after cleaning", len(cleaned))
    return cleaned


# ──────────────────────────── Format-Specific ────────────────────────


def parse_pdf(file: Union[str, Path, io.BytesIO]) -> str:
    """Extract text from every page of a PDF."""
    pages: list[str] = []
    with pdfplumber.open(file) as pdf:
        for i, page in enumerate(pdf.pages):
            text = page.extract_text()
            if text:
                pages.append(text)
            else:
                logger.debug("Page %d yielded no text (image-only?)", i + 1)
    return "\n\n".join(pages)


def parse_docx(file: Union[str, Path, io.BytesIO]) -> str:
    """Extract text from paragraphs and tables of a DOCX file."""
    doc = DocxDocument(file)
    parts: list[str] = []

    # Paragraphs
    for para in doc.paragraphs:
        if para.text.strip():
            parts.append(para.text)

    # Tables (flatten cells row-by-row)
    for table in doc.tables:
        for row in table.rows:
            row_text = " | ".join(cell.text.strip() for cell in row.cells if cell.text.strip())
            if row_text:
                parts.append(row_text)

    return "\n\n".join(parts)


# ──────────────────────────── Helpers ─────────────────────────────────


def _detect_extension(file: Union[str, Path, io.BytesIO], filename: str) -> str:
    """Return normalised file extension (e.g. '.pdf')."""
    if isinstance(file, (str, Path)):
        return Path(file).suffix.lower()
    if filename:
        return Path(filename).suffix.lower()
    raise ValueError("Cannot detect file format — provide a filename.")
